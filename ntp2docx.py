#from md2py import md2py
import fnmatch
import os
import re
import mistune
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

# following function from https://stackoverflow.com/a/48666968/1018206
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

# following function modified from https://github.com/python-openxml/python-docx/issues/846#issuecomment-660494075
def add_bookmark(paragraph, _id):
    """Insert bookmark to a paragraph."""
    name = '_Ref_id_num_' + _id
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:name'), name)
    start.set(qn('w:id'), str(_id))
    paragraph._p.append(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(_id))
    paragraph._p.append(end)
    return name

# following function modified from https://github.com/python-openxml/python-docx/issues/846#issuecomment-660494075
def append_ref_to_paragraph(paragraph, refName, text = ''):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), refName)
    hyperlinktext = OxmlElement('w:r')
    hyperlinktextproperty = OxmlElement('w:rPr')
    hyperlinktextpropertystyle = OxmlElement('w:rStyle')
    hyperlinktextpropertystyle.set(qn('w:val'), "InternetLink")
    hyperlinktextcontent = OxmlElement('w:t')
    hyperlinktextcontent.text = text
    hyperlinktextproperty.append(hyperlinktextpropertystyle)
    hyperlinktext.append(hyperlinktextproperty)
    hyperlinktext.append(hyperlinktextcontent)
    hyperlink.append(hyperlinktext)
    paragraph._p.append(hyperlink)

def find(pattern, path):
    """Find files in a path based on a filename pattern with numbers and sorted by the number ASC."""
    result = {}
    for root, dirs, files in os.walk(path):
        for name in files:
            if fnmatch.fnmatch(name, pattern):
                tpfile = os.path.join(root, name)
                tpnumber = re.findall("[0-9]+", tpfile)
                result[tpfile] = int(tpnumber[0])
    result = sorted(result.items(), key=lambda kv: kv[1])
    return result

# ingest group of markdown files describing properties into the main document
def ingest(files, curpara):
    properties = {}

    for file in files:
        mdfile = open(file[0], "r").read()
        html = mistune.markdown(mdfile)
        parsed_html = BeautifulSoup(html, features="lxml")
        heading1 = parsed_html.find_all('h1')[0].contents[0]  # should be only one
        propid = str.lower(heading1[:heading1.find(" ")])
        curpara = insert_paragraph_after(curpara, heading1, "CRM Property Label")
        curbookmark = add_bookmark(curpara, propid)
        properties[propid] = [curpara, curbookmark, parsed_html]

    for propertyname, property in properties.items():
        headings2 = property[2].find_all('h2')
        curpara = property[0]
        for heading2 in headings2:
            curpara = insert_paragraph_after(curpara, heading2.contents[0], "CRM Description Label")
            if heading2.contents[0] == "Examples:":
                curtextruns = heading2.find_next("ul")
            else:
                curtextruns = heading2.find_next("p")

            if heading2.contents[0] == "Domain:" or heading2.contents[0] == "Range:":
                curstyle = "CRM Domain Range"
            elif heading2.contents[0] == "Superproperty of:":
                curstyle = "CRM Super Sub Property"
            elif heading2.contents[0] == "Subproperty of:":
                curstyle = "CRM Super Sub Property"
            elif heading2.contents[0] == "Quantification:":
                curstyle = "CRM Quantification"
            elif heading2.contents[0] == "Scope note:":
                curstyle = "CRM Scope Note Text"
            elif heading2.contents[0] == "Examples:":
                curstyle = "CRM Example"
            elif heading2.contents[0] == "In First Order Logic:":
                curstyle = "CRM First Order Logic"

            curpara = insert_paragraph_after(curpara, None, curstyle)
            if curtextruns.name == 'ul': # if this is a list
                curtextruns.contents = curtextruns.contents[1:-1] # remove leading and ending linefeeds
            for curtextrun in curtextruns:
                if curtextrun.name == "a":  # if this part of the run is a link
                    for attr, curlink in curtextrun.attrs.items():
                        if attr == 'href':
                            curlink = "_Ref_id_num_" + curlink[1:]  # trim the '#' from markdown
                            append_ref_to_paragraph(curpara, curlink, curtextrun.contents[0])
                elif curtextrun.name == "li": # if this is a list of examples
                    curpara.add_run(curtextrun.contents[0])
                elif curtextrun.name == "em": # if this is italics
                    runner = curpara.add_run(curtextrun.contents[0])
                    runner.italic = True
                elif curtextrun.name == "strong": # if this is bold
                    runner = curpara.add_run(curtextrun.contents[0])
                    runner.bold = True
                else:  # treat everything else as text for now
                    curpara.add_run(curtextrun)
    return curpara

# open template.docx
document = Document("template.docx")

# find the heading of the property declarations
pdparagraphs = document.paragraphs
for pdparagraph in pdparagraphs:
    if pdparagraph.text == "CRMntp Property Declarations":
         break

# find all files with tp property definitions as included in the repo
files = find('tp*.md', 'source')
# ingest them and return the last paragraph of the ingestion
lasttpparagraph = ingest(files, pdparagraph)
# find all files with ntp property definitions as included in the repo
files = find('ntp*.md', 'source')
# ingest them and return the last paragraph of the ingestion
lastntpparagraph = ingest(files, lasttpparagraph)

# # add table of contents from here: https://github.com/python-openxml/python-docx/issues/36#issuecomment-145302669
# # but it works on Windows MS Word only I am afraid...
# # find the heading of the index
# indexparagraphs = document.paragraphs
# for indexparagraph in indexparagraphs:
#     if indexparagraph.text == "Index":
#          break
#
# paragraph = insert_paragraph_after(indexparagraph)
# run = paragraph.add_run()
# fldChar = OxmlElement('w:fldChar')  # creates a new element
# fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
# instrText = OxmlElement('w:instrText')
# instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
# instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need
#
# fldChar2 = OxmlElement('w:fldChar')
# fldChar2.set(qn('w:fldCharType'), 'separate')
# fldChar3 = OxmlElement('w:t')
# fldChar3.text = "Right-click to update field."
# fldChar2.append(fldChar3)
#
# fldChar4 = OxmlElement('w:fldChar')
# fldChar4.set(qn('w:fldCharType'), 'end')
#
# r_element = run._r
# r_element.append(fldChar)
# r_element.append(instrText)
# r_element.append(fldChar2)
# r_element.append(fldChar4)
# p_element = paragraph._p

document.save('demo.docx')